"""DOCX renderer.

Renders the document design from the content tree. Everything that was
hand-maintained in the legacy generators is derived here: heading numbers, the
table of contents, and figure numbers all come from the outline, so inserting a
section can never leave the document internally inconsistent.
"""
from __future__ import annotations

import os
import re
import tempfile
from dataclasses import replace
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

from ..imaging import auto_crop_image
from ..theme import THEME, Theme

ICON_MARKER = re.compile(r"\[icon:([^\]\s]+)(?:\s+=([\d.]+)cm)?\]")

# One millimetre in twips, which is what Word measures a tab stop in.
TWIPS_PER_MM = 56.6929


def logo_path(project) -> Path | None:
    """The company's mark for the cover, or None if this document names none.

    Set as ``document.logo`` in content/doc.yaml, a path relative to the project
    root so what a person types is what they see in their own folder. The asset
    store is tried as well, because that is where everything else a document
    draws already lives.

    A mark that is named but not on disk is left out rather than fatal. A cover
    without a logo still prints, and stopping a release over one file helps
    nobody; `verba lint` is where a person should be told.

    It lives in this module because the other two renderers already import from
    here, so all three reach it without a circular import.
    """
    name = str((project.config.get("document") or {}).get("logo") or "").strip()
    if not name:
        return None
    candidate = Path(name)
    if not candidate.is_absolute():
        candidate = Path(getattr(project, "root", ".")) / candidate
    if candidate.exists():
        return candidate
    try:
        in_assets = project.asset_path(name)
    except Exception:
        return None
    return in_assets if in_assets.exists() else None


def C(h: str) -> RGBColor:
    return RGBColor.from_string(h)


# ── low level OOXML helpers ──────────────────────────────────────────────────

def _el(tag: str, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(f"w:{k}"), v)
    return e


def set_shd(para, fill: str):
    para._p.get_or_add_pPr().append(_el("w:shd", val="clear", color="auto", fill=fill))


def set_left_border(para, color: str, sz="24", space="8"):
    pbdr = _el("w:pBdr")
    pbdr.append(_el("w:left", val="single", sz=sz, space=space, color=color))
    para._p.get_or_add_pPr().append(pbdr)


def set_edge_border(para, edge: str, color: str, sz="8"):
    pbdr = _el("w:pBdr")
    pbdr.append(_el(f"w:{edge}", val="single", sz=sz, space="4", color=color))
    para._p.get_or_add_pPr().append(pbdr)


def set_indent(para, left_twips=280, hanging=0):
    ind = _el("w:ind", left=str(left_twips))
    if hanging:
        ind.set(qn("w:hanging"), str(hanging))
    para._p.get_or_add_pPr().append(ind)


def add_right_tab(para, pos=9072):
    tabs = _el("w:tabs")
    tabs.append(_el("w:tab", val="right", pos=str(pos)))
    para._p.get_or_add_pPr().append(tabs)


def add_page_field(para):
    r = para.add_run()
    for tag, attrs, text in (("w:fldChar", {"fldCharType": "begin"}, None),
                             ("w:instrText", {"xml:space": "preserve"}, " PAGE "),
                             ("w:fldChar", {"fldCharType": "end"}, None)):
        e = OxmlElement(tag)
        for k, v in attrs.items():
            e.set(qn(k) if ":" in k else qn(f"w:{k}"), v)
        if text:
            e.text = text
        r._r.append(e)
    return r


# style_run is module level and predates the font being a choice. Rather than
# thread the renderer through every call site, the active family is published
# here when a render begins.
_ACTIVE_FONT = [THEME.font]


def style_run(run, size=None, bold=False, italic=False, color=None, font=None):
    run.font.name = font or _ACTIVE_FONT[0]
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = C(color)
    return run


# ── renderer ─────────────────────────────────────────────────────────────────


class DocxRenderer:
    def __init__(self, project, theme=None):
        self.p = project
        # The theme is the project's, not a module-level default: the engine
        # documents whatever product it is pointed at, and each one has its own.
        if theme is None:
            theme = Theme.load(getattr(project, "root", "."))
        # The chosen typeface travels through the theme, so every existing
        # `self.t.font` reference picks it up without being rewritten.
        from ..typography import Typography
        typo = Typography.load(getattr(project, "root", "."))
        face = typo.face("document")
        self.face = face
        # The sheet and the margins come from the same setting the PDF reads,
        # so the two outputs cannot be laid out differently. They were: this
        # renderer had A4 at 25mm written into it while the PDF took 18mm from
        # content/typography.yaml, and nothing said which one was the document.
        self.page = typo.page
        self.t = replace(theme, font=face.docx)
        self.doc = Document()
        self.used_images: dict[str, str] = {}      # asset -> first section that used it
        self.duplicate_uses: list[tuple[str, str, str]] = []
        self.last_image: str | None = None
        self.consecutive_repeats: list[str] = []
        self.figure_counter: dict[str, int] = {}
        self._tmp: list[str] = []

    # -- document chrome ----------------------------------------------------
    def _setup(self):
        doc = self.doc
        sec = doc.sections[0]
        pg = self.page
        sec.page_width, sec.page_height = Mm(pg.width_mm), Mm(pg.height_mm)
        sec.left_margin = sec.right_margin = Mm(pg.side)
        sec.top_margin, sec.bottom_margin = Mm(pg.margin_top), Mm(pg.margin_bottom)
        sec.different_first_page_header_footer = True

        for name, bold, size, color, before, after in [
            ("Heading 3", True, self.t.size_h3, self.t.navy_deep, 14, 4),
            ("Heading 4", True, self.t.size_h4, self.t.brand_blue, 10, 3),
            ("Normal", False, self.t.size_body, self.t.navy_deep, 0, 5),
        ]:
            s = doc.styles[name]
            s.font.name = self.t.font
            _rfonts(s, self.t.font)
            s.font.size = Pt(size)
            s.font.bold = bold
            s.font.color.rgb = C(color)
            s.paragraph_format.space_before = Pt(before)
            s.paragraph_format.space_after = Pt(after)

        # Word only honours a family it can find. Naming the fallback in the
        # same run properties means a reader without the face gets a considered
        # substitute rather than whatever Word decides on its own.
        _ACTIVE_FONT[0] = self.t.font
        _rfonts(doc.styles["Normal"], self.t.font, self.face.docx_fallback)

        self._header_footer(sec)
        cp = doc.core_properties
        cp.title = f"{self.p.config['product']['name']} {self.p.title()}"
        cp.author = (self.p.config.get("document", {}).get("author")
                     or self.p.config.get("product", {}).get("vendor", ""))
        cp.comments = f"Built by verba, profile={self.p.profile.name}"

    def _header_footer(self, section):
        prod = self.p.config["product"]["name"]
        para = section.header.paragraphs[0]
        para.clear()
        add_right_tab(para)
        set_edge_border(para, "bottom", self.t.brand_blue, sz="12")
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(6)
        # The running header names the product being documented, not the tool
        # that built it, and not the company that wrote the tool.
        vendor = self.p.config.get("product", {}).get("vendor", "")
        mark = f"{vendor.upper()}  |  {prod}" if vendor and vendor != prod else prod
        style_run(para.add_run(mark), self.t.size_chrome + 0.5, True,
                  color=self.t.brand_blue)
        style_run(para.add_run("\t"), self.t.size_chrome + 0.5)
        style_run(para.add_run("Page "), self.t.size_chrome + 0.5, color=self.t.grey_mid)
        style_run(add_page_field(para), self.t.size_chrome + 0.5, color=self.t.grey_mid)

        para = section.footer.paragraphs[0]
        para.clear()
        add_right_tab(para)
        set_edge_border(para, "top", self.t.grey_mid, sz="4")
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(0)
        style_run(para.add_run(f"{prod}: {self.p.title()}"), self.t.size_chrome,
                  color=self.t.grey_mid)
        conf = self.p.config.get("document", {}).get("confidentiality", "")
        stamp = self.p.config.get("_release_label") or date.today().strftime("%B %Y")
        style_run(para.add_run(f"\t{conf}  |  {stamp}"), self.t.size_chrome,
                  color=self.t.grey_mid)

    # -- headings -----------------------------------------------------------
    def h1(self, text):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before, pf.space_after = Pt(0), Pt(10)
        pf.keep_with_next = True
        pf.page_break_before = True
        style_run(p.add_run(text), self.t.size_h1, True, color=self.t.navy_deep)
        return p

    def h2(self, text):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before, pf.space_after = Pt(22), Pt(5)
        pf.keep_with_next = True
        style_run(p.add_run(text), self.t.size_h2, True, color=self.t.brand_blue)
        return p

    def h3(self, text):
        p = self.doc.add_heading(text, level=3)
        p.paragraph_format.keep_with_next = True
        return p

    def h4(self, text):
        p = self.doc.add_heading(text, level=4)
        p.paragraph_format.keep_with_next = True
        return p

    # -- blocks -------------------------------------------------------------
    def body(self, text, keep=False):
        p = self.doc.add_paragraph()
        style_run(p.add_run(text), self.t.size_body, color=self.t.navy_deep)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.keep_with_next = keep
        return p

    def bullets(self, items):
        for item in items:
            if not item:
                continue
            p = self.doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            self._runs_with_icons(p, str(item), self.t.size_body, self.t.navy_deep)

    def steps(self, items):
        for n, item in enumerate(items, 1):
            p = self.doc.add_paragraph()
            set_indent(p, left_twips=340, hanging=340)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            style_run(p.add_run(f"{n}.  "), self.t.size_body, True, color=self.t.brand_blue)
            self._runs_with_icons(p, str(item), self.t.size_body, self.t.navy_deep)

    def _runs_with_icons(self, para, text, size, color):
        """Split ``[icon:name.png]`` markers out of a line into inline images."""
        pos = 0
        for m in ICON_MARKER.finditer(text):
            head = text[pos:m.start()].rstrip()
            if head:
                style_run(para.add_run(head), size, color=color)
            name, width = m.group(1), m.group(2)
            path = self.p.asset_path(name)
            if path.exists():
                try:
                    img = Image.open(path)
                    w_cm = float(width) if width else min(4.2, img.size[0] / 58.0)
                    para.add_run(" ").font.size = Pt(size)
                    para.add_run().add_picture(str(path), width=Cm(w_cm))
                    self._note_image_use(name)
                except Exception:
                    style_run(para.add_run(f"[{name}]"), size, italic=True,
                              color=self.t.grey_mid)
            pos = m.end()
        tail = text[pos:]
        if tail.strip():
            style_run(para.add_run(tail if pos == 0 else tail.lstrip()), size, color=color)

    def note(self, text, label="Note"):
        # note_icons holds the NAME of a drawn mark, not a character. The HTML
        # and PDF paths draw the SVG; this one concatenated the name into a run,
        # so every callout in the shipped Word file read "note Note:" and
        # "warning Important:". Ten of them are in the released v31. Word has no
        # SVG in a run, so the label carries the callout on its own.
        icon = ""
        accent = self.t.note_accent.get(label, self.t.brand_blue)
        p = self.doc.add_paragraph()
        set_shd(p, self.t.lavender)
        set_left_border(p, accent, sz="16", space="6")
        set_indent(p, left_twips=200)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        style_run(p.add_run(f"{icon}{label}:  ".lstrip()),
                  self.t.size_small + 0.5, True, color=accent)
        style_run(p.add_run(text), self.t.size_small + 0.5, color=self.t.navy_deep)

    def fields(self, items):
        for f in items:
            p = self.doc.add_paragraph()
            if f.get("required"):
                style_run(p.add_run("* "), self.t.size_small, True, color=self.t.red_err)
            style_run(p.add_run(f.get("field", "")), self.t.size_body, True,
                      color=self.t.navy_deep)
            if f.get("type"):
                style_run(p.add_run(f"  {f['type']}"), self.t.size_small,
                          color=self.t.periwinkle)
            if f.get("description"):
                style_run(p.add_run(":  "), self.t.size_small, color=self.t.grey_mid)
                style_run(p.add_run(f["description"]), self.t.size_body,
                          color=self.t.navy_deep)
            set_indent(p, left_twips=280, hanging=280)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)

    def _labelled(self, items, key, color, space=3):
        for it in items:
            p = self.doc.add_paragraph()
            style_run(p.add_run(it.get(key, "")), self.t.size_body, True, color=color)
            if it.get("description"):
                style_run(p.add_run(":  "), self.t.size_small, color=self.t.grey_mid)
                style_run(p.add_run(it["description"]), self.t.size_body,
                          color=self.t.navy_deep)
            set_indent(p, left_twips=280, hanging=280)
            p.paragraph_format.space_before = Pt(space)
            p.paragraph_format.space_after = Pt(space)

    def actions(self, items):
        self._labelled(items, "action", self.t.brand_blue)

    def columns(self, items, key: str = "column"):
        self._labelled(items, key, self.t.brand_blue, space=2)

    def terms(self, items):
        for it in items:
            p = self.doc.add_paragraph()
            style_run(p.add_run(f"{it.get('term','')}:  "), self.t.size_body, True,
                      color=self.t.navy_deep)
            style_run(p.add_run(it.get("definition", "")), self.t.size_body,
                      color=self.t.grey_dark)
            set_indent(p, left_twips=280, hanging=280)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)

    def _note_image_use(self, name):
        if name in self.used_images:
            self.duplicate_uses.append((name, self.used_images[name], self.current_section))
        else:
            self.used_images[name] = self.current_section

    def screenshot(self, attrs, chapter_number):
        name = attrs.get("file", "")
        path = self.p.asset_path(name)
        if not path.exists():
            p = self.doc.add_paragraph()
            style_run(p.add_run(f"[missing asset: {name}]"), self.t.size_body,
                      italic=True, color=self.t.red_err)
            return
        if name == self.last_image:
            self.consecutive_repeats.append(name)
        self.last_image = name
        self._note_image_use(name)

        width = float(attrs.get("width_cm") or
                      self.p.config.get("build", {}).get("screenshot_width_cm", 15.0))
        img = auto_crop_image(Image.open(path))
        tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        img.save(tmp.name, "PNG")
        tmp.close()
        self._tmp.append(tmp.name)
        self.doc.add_picture(tmp.name, width=Cm(width))
        self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        caption = attrs.get("caption") or ""
        self.figure_counter[chapter_number] = self.figure_counter.get(chapter_number, 0) + 1
        label = f"Figure {chapter_number}.{self.figure_counter[chapter_number]}"
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        style_run(cap.add_run(f"{label}"), self.t.size_caption, True, color=self.t.brand_blue)
        if caption:
            style_run(cap.add_run(f": {caption}"), self.t.size_caption, italic=True,
                      color=self.t.grey_mid)

    # -- front matter --------------------------------------------------------
    def cover(self):
        cfg, prod = self.p.config, self.p.config["product"]
        doc_cfg = cfg.get("document", {})
        # A label with nothing beside it reads as a field that failed to fill,
        # not as a fact that does not apply. The PDF cover has always dropped
        # these; this one printed "Environment:" against blank space.
        rows = [(k, v) for k, v in (
            ("Environment", doc_cfg.get("environment", "")),
            ("Platform", prod.get("platform_version", "")),
            ("Edition", self.p.profile.name.title()),
            ("Revision", cfg.get("_release_label", "draft")),
        ) if str(v).strip()]
        tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        _render_cover(tmp.name, prod.get("vendor", ""), prod.get("name", ""),
                      self.p.title(), rows, doc_cfg.get("confidentiality", ""),
                      theme=self.t, logo=logo_path(self.p))
        tmp.close()
        self._tmp.append(tmp.name)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(tmp.name, width=Cm(16))
        self.doc.add_page_break()

    def toc(self):
        self.h1("Table of Contents")
        depth = self.p.config.get("build", {}).get("toc_depth", 3)
        for node in self.p.nodes:
            if node.level > depth:
                continue
            p = self.doc.add_paragraph()
            set_indent(p, left_twips=(node.level - 1) * 320)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            top = node.level == 1
            style_run(p.add_run(f"{node.number}  "), self.t.size_body, top,
                      color=self.t.brand_blue if top else self.t.grey_mid)
            style_run(p.add_run(node.title), self.t.size_body, top,
                      color=self.t.navy_deep if top else self.t.grey_dark)
        self.doc.add_page_break()

    def revision_history(self, history):
        if not history:
            return
        self.h1("Revision History")
        self.body("Each revision below was produced from the versioned content "
                  "sources and rebuilt end to end. Entries are newest first.")
        for entry in history:
            p = self.doc.add_paragraph()
            set_indent(p, left_twips=280, hanging=280)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            style_run(p.add_run(f"{entry.get('version','')}  "), self.t.size_body, True,
                      color=self.t.brand_blue)
            style_run(p.add_run(f"{entry.get('date','')}  "), self.t.size_small,
                      color=self.t.grey_mid)
            style_run(p.add_run(entry.get("summary", "")), self.t.size_body,
                      color=self.t.navy_deep)

    # -- main ---------------------------------------------------------------
    def render(self, out_path: Path, history=None) -> Path:
        self._setup()
        self.cover()
        self.toc()
        if history:
            self.revision_history(history)

        for node in self.p.nodes:
            sec = node.section
            if sec is None:
                continue
            self.current_section = node.id
            chapter = node.number.split(".")[0]
            title = f"{node.number}. {sec.title}" if node.level == 1 else \
                    f"{node.number} {sec.title}"
            if sec.icon:
                title = f"{sec.icon} {title}"
            {1: self.h1, 2: self.h2, 3: self.h3}.get(node.level, self.h4)(title)

            first_body = True
            for block in self.p.resolved_blocks(sec):
                k = block.kind
                if k == "paragraph":
                    self.body(block.text, keep=(first_body and node.level >= 2))
                    first_body = False
                elif k == "bullets":
                    self.bullets(block.items)
                elif k == "steps":
                    self.steps(block.items)
                elif k == "fields":
                    self.fields(block.items)
                elif k == "actions":
                    self.actions(block.items)
                elif k == "columns":
                    self.columns(block.items)
                elif k == "tabs":
                    self.columns(block.items, key="tab")
                elif k == "terms":
                    self.terms(block.items)
                elif k == "note":
                    self.note(block.text, block.attrs.get("label", "Note"))
                elif k == "screenshot":
                    self.screenshot(block.attrs, chapter)
                elif k == "heading":
                    self.h4(block.text)
                    self.last_image = None

        out_path = Path(out_path)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(out_path))
        for t in self._tmp:
            try:
                os.unlink(t)
            except OSError:
                pass
        return out_path


def _rfonts(style, family: str, fallback: str | None = None):
    """Write w:rFonts so the family applies to every script Word distinguishes.

    python-docx sets only the ASCII slot. A document whose headings survive but
    whose tables come out in Cambria is this, every time.
    """
    rpr = style.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    for slot in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(slot), family)
    if fallback:
        rf.set(qn("w:hint"), "default")


# ── cover art ────────────────────────────────────────────────────────────────

# Tried in order, and Pillow's own DejaVu is last because it ships with the
# package and therefore always exists. Two absolute macOS paths were the whole
# list, so on Linux and Windows every string on this cover fell back to
# Pillow's default bitmap face and the page came out effectively blank. The
# package is published for everybody.
_BOLD_FACES = ("/System/Library/Fonts/Supplemental/Arial Bold.ttf",
               "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
               "C:/Windows/Fonts/arialbd.ttf",
               "DejaVuSans-Bold.ttf")
_REG_FACES = ("/System/Library/Fonts/Supplemental/Arial.ttf",
              "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
              "C:/Windows/Fonts/arial.ttf",
              "DejaVuSans.ttf")
_FONT_BOLD = _BOLD_FACES
_FONT_REG = _REG_FACES


def _font(faces, size):
    """The first of these that this machine actually has."""
    for face in (faces if isinstance(faces, (tuple, list)) else [faces]):
        try:
            return ImageFont.truetype(face, size)
        except Exception:
            continue
    return ImageFont.load_default()


def _fits(draw, text, faces, start, width):
    """The largest size at which this string fits, down to a floor.

    The vendor size was tuned for the four letters of RISE and nothing measured
    the string, so any company with a longer name had it cut off at the edge of
    its own cover.
    """
    size = start
    while size > 28:
        font = _font(faces, size)
        box = draw.textbbox((0, 0), text, font=font)
        if box[2] - box[0] <= width:
            return font
        size -= 4
    return _font(faces, 28)


def _rgb(value, fallback):
    """A theme token as an (r, g, b) triple."""
    v = str(value or "").strip().lstrip("#")
    if len(v) == 6:
        try:
            return tuple(int(v[i:i + 2], 16) for i in (0, 2, 4))
        except ValueError:
            pass
    return fallback


def _render_cover(out, vendor, product, subtitle, rows, confidentiality,
                  W=1240, H=1876, theme=None, logo=None):
    # The palette is the document's, not this file's. These five literals meant
    # that whichever theme a person chose, and whatever palette they authored,
    # the DOCX cover came out in one company's indigo while the body of the
    # same file followed their colours.
    theme = theme or THEME
    navy = _rgb(getattr(theme, "navy_deep", None), (27, 37, 73))
    blue = _rgb(getattr(theme, "brand_blue", None), (49, 55, 219))
    lav = _rgb(getattr(theme, "lavender", None), (235, 239, 252))
    grey, greydk = (161, 161, 161), (100, 100, 110)
    img = Image.new("RGB", (W, H), (255, 255, 255))
    d = ImageDraw.Draw(img)
    d.rectangle([(0, 0), (W, 30)], fill=blue)

    top = 440
    if logo is not None:
        try:
            mark = Image.open(logo).convert("RGBA")
            wide = 320
            mark.thumbnail((wide, 200))
            img.paste(mark, (80, 250), mark)
            top = 250 + mark.height + 60
        except Exception:
            # A mark that cannot be opened is left out. A cover without a logo
            # still prints, and failing a release over one file helps nobody.
            pass

    d.text((80, top), vendor.upper(),
           font=_fits(d, vendor.upper(), _FONT_BOLD, 130, W - 160), fill=navy)
    rule = top + 170
    d.rectangle([(80, rule), (W - 80, rule + 4)], fill=blue)
    d.text((82, rule + 22),
           product, font=_fits(d, product, _FONT_REG, 66, W - 164), fill=navy)
    d.text((84, rule + 116), subtitle,
           font=_fits(d, subtitle, _FONT_REG, 27, W - 168), fill=blue)
    y1, y2 = 950, 950 + 40 + len(rows) * 54
    d.rounded_rectangle([(80, y1), (W - 80, y2)], radius=10, fill=lav)
    d.rectangle([(80, y1), (85, y2)], fill=blue)
    for i, (label, value) in enumerate(rows):
        y = y1 + 28 + i * 54
        d.text((108, y), f"{label}:", font=_font(_FONT_BOLD, 22), fill=navy)
        d.text((290, y), str(value), font=_font(_FONT_REG, 22), fill=greydk)
    bbox = d.textbbox((0, 0), confidentiality, font=_font(_FONT_REG, 19))
    d.text(((W - (bbox[2] - bbox[0])) // 2, H - 68), confidentiality,
           font=_font(_FONT_REG, 19), fill=grey)
    d.rectangle([(0, H - 30), (W, H)], fill=blue)
    img.save(out, dpi=(150, 150))
